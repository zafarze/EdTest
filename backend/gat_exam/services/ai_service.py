import os
import json
import base64
from pathlib import Path
from openai import OpenAI
from dotenv import load_dotenv
from pypdf import PdfReader
import docx
from django.db.models import Avg

# 🔥 Импортируем наш новый мозг (Brain Center)
# Поскольку оба файла лежат в папке services, используем относительный импорт
from .prompt_service import PromptService

# --- 1. НАСТРОЙКА ОКРУЖЕНИЯ ---
BASE_DIR = Path(__file__).resolve().parent.parent.parent.parent
# Корректировка пути зависит от вложенности, но обычно .env в корне
env_path = os.path.join(BASE_DIR, '.env')

load_dotenv(env_path)
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)


# -------------------------------------------------------------------------
# 2. ПАРСИНГ ФАЙЛОВ (Теперь тоже через PromptService!)
# -------------------------------------------------------------------------
def parse_file_with_ai(file_obj, filename):
    print(f"\n📂 [AI Service] Начало обработки файла: {filename}")
    if not client.api_key:
        print("❌ Ошибка: API ключ OpenAI не найден.")
        return []

    text_content = ""
    image_content = None
    ext = filename.split('.')[-1].lower()

    try:
        if ext == 'pdf':
            reader = PdfReader(file_obj)
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
        elif ext in ['docx', 'doc']:
            doc = docx.Document(file_obj)
            para_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            text_content += para_text + "\n"
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data: table_text += " | ".join(row_data) + "\n"
            text_content += table_text
        elif ext in ['jpg', 'jpeg', 'png', 'webp']:
            file_obj.seek(0)
            encoded_image = base64.b64encode(file_obj.read()).decode('utf-8')
            image_content = {
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{encoded_image}"}
            }
        else:
            return []
    except Exception as e:
        print(f"❌ Ошибка чтения файла: {e}")
        return []

    if not text_content.strip() and not image_content: return []
    
    # 🔥 ИСПОЛЬЗУЕМ PROMPT SERVICE (Slug: file_parser)
    # Контекст передаем пустым, так как контент добавляем динамически ниже
    messages, config = PromptService.format_messages("file_parser", {})
    
    # Добавляем контент файла вручную в сообщение пользователя
    user_payload = []
    if text_content: 
        user_payload.append({"type": "text", "text": f"Текст файла:\n{text_content[:25000]}"})
    if image_content: 
        user_payload.append(image_content)
    
    # Перезаписываем user message, так как контент сложный (текст + фото)
    # messages[0] - это system, messages[1] - это user (шаблон из PromptService)
    # Мы заменяем контент user message на реальный файл
    messages[1]["content"] = user_payload

    try:
        response = client.chat.completions.create(
            model=config['model'],
            messages=messages,
            response_format={"type": "json_object"},
            temperature=config['temp']
        )
        data = json.loads(response.choices[0].message.content)
        if "questions" in data: return data["questions"]
        elif isinstance(data, list): return data
        return []
    except Exception as e:
        print(f"❌ AI Parsing Error: {e}")
        return []


# -------------------------------------------------------------------------
# 🔥 3. АНАЛИЗ ВОПРОСА (СУПЕР-МОЗГ + ЗРЕНИЕ)
# -------------------------------------------------------------------------
def analyze_question_ai(text, choices, image_file=None):
    """
    Проверяет: Факты, Зрение, Грамматику.
    Логика полностью вынесена в PromptService (Slug: question_audit).
    """
    print(f"\n🧐 [AI AUDIT] Проверка: '{text[:30]}...'")

    if not client.api_key:
        return {"valid": True, "message": "API Key not found."}

    # 1. Готовим текст вариантов
    choices_str = "\n".join([
        f"- {c.get('text', '')} {'(CORRECT)' if c.get('is_correct') else ''}" 
        for c in choices
    ])

    # 2. Формируем контекст для шаблона
    context = {
        "text": text,
        "choices": choices_str
    }

    # 3. Получаем сообщения из "Мозгового Центра"
    messages, config = PromptService.format_messages("question_audit", context)

    # 4. Обработка картинки (если есть)
    # Картинку нельзя вставить в текстовый шаблон {image}, поэтому добавляем её нативно в API
    if image_file:
        try:
            image_file.seek(0)
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            image_payload = {
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
            }
            
            # user message находится под индексом 1
            # Если там просто строка (от шаблона), превращаем её в список content
            current_content = messages[1]["content"]
            if isinstance(current_content, str):
                messages[1]["content"] = [
                    {"type": "text", "text": current_content},
                    image_payload
                ]
            elif isinstance(current_content, list):
                messages[1]["content"].append(image_payload)
                
            print("📸 Картинка добавлена к анализу")
        except Exception as e:
            print(f"⚠️ Ошибка чтения картинки: {e}")

    # 5. Вызов API
    try:
        response = client.chat.completions.create(
            model=config['model'], # GPT-4o или что настроено в админке
            messages=messages,
            response_format={"type": "json_object"},
            temperature=config['temp'] # Креативность из админки
        )
        
        result = json.loads(response.choices[0].message.content)
        print(f"✅ [AI RESULT]: {result}")
        return result

    except Exception as e:
        print(f"❌ AI Error: {e}")
        return {"valid": True, "message": f"Сбой проверки AI: {str(e)}"}


# -------------------------------------------------------------------------
# 4. ГЕНЕРАЦИЯ ДИСТРАКТОРОВ
# -------------------------------------------------------------------------
def generate_distractors_ai(question_text, correct_answer):
    if not client.api_key: return ["Error", "No", "Key"]
    
    # 1. Контекст
    context = {
        "text": question_text,
        "answer": correct_answer
    }
    
    # 2. Получаем промпт (Slug: distractor_gen)
    messages, config = PromptService.format_messages("distractor_gen", context)

    try:
        response = client.chat.completions.create(
            model=config['model'],
            messages=messages,
            response_format={"type": "json_object"},
            temperature=config['temp']
        )
        data = json.loads(response.choices[0].message.content)
        return data.get("distractors", [])[:3]
    except Exception as e:
        print(f"Distractor Gen Error: {e}")
        return ["Ошибка", "Генерации", "AI"]


# -------------------------------------------------------------------------
# 5. ОТЧЕТ ПО КЛАССУ
# -------------------------------------------------------------------------
def generate_class_report(exam_id):
    """
    Анализирует результаты и пишет отчет учителю.
    """
    # Ленивый импорт моделей
    from ..models import Exam, ExamResult
    
    if not client.api_key: return "Ошибка: AI ключ не настроен."

    try:
        exam = Exam.objects.get(id=exam_id)
        results = ExamResult.objects.filter(exam=exam)
        total_students = results.count()

        if total_students < 1:
            return "Нет данных для анализа."

        # Статистика
        avg_score_val = results.aggregate(Avg('percentage'))['percentage__avg'] or 0
        
        exam_topic = "Общий экзамен"
        if exam.questions.exists() and exam.questions.first().topic:
            exam_topic = exam.questions.first().topic.title

        # 1. Контекст
        context = {
            "topic": exam_topic,
            "count": total_students,
            "avg": f"{avg_score_val:.1f}"
        }

        # 2. Получаем промпт (Slug: class_report)
        # Если в базе нет промпта 'class_report', добавь его в DEFAULTS PromptService!
        messages, config = PromptService.format_messages("class_report", context)

        response = client.chat.completions.create(
            model=config['model'],
            messages=messages,
            temperature=config['temp']
        )
        
        return response.choices[0].message.content

    except Exception as e:
        print(f"Report Error: {e}")
        return f"Ошибка при генерации отчета: {str(e)}"